import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

st.set_page_config(page_title="Generador de Minuta FAF", layout="centered")

st.title("📝 Generador de Minuta FAF")

st.markdown("Completa los datos para generar la minuta:")

# Datos generales
fecha = st.text_input("📅 Fecha de la sesión/cooperación", "")
meses = st.text_input("🗓️ Meses que abarca la sesión", "")
ambigu = st.text_input("🍽️ ¿Quién se encargó del ambigú?", "")

st.subheader("📥 Ingresos")
ingresos = {
    "Ingresos": st.number_input("Ingresos", min_value=0.0),
    "Ahorros": st.number_input("Ahorros", min_value=0.0),
    "Intereses por préstamos": st.number_input("Intereses por préstamos", min_value=0.0),
    "Multas acumuladas": st.number_input("Multas acumuladas", min_value=0.0),
    "Aportación del patrimonio del FAF hasta el año 2025": st.number_input("Aportación del patrimonio", min_value=0.0),
}

st.subheader("📤 Egresos")
egresos = {
    "Egresos": st.number_input("Egresos", min_value=0.0),
    "Préstamos": st.number_input("Préstamos", min_value=0.0),
    "Efectivo en caja": st.number_input("Efectivo en caja", min_value=0.0),
    "Citibanamex": st.number_input("Citibanamex", min_value=0.0),
    "Banorte": st.number_input("Banorte", min_value=0.0),
    "Inversión en banco": st.number_input("Inversión en banco", min_value=0.0),
}

miembros = [
    "Alfonso Rodríguez Grado", "Andres Chaparro Campuzano", "Andres Chaparro Montoya",
    "Angelica Cervantes", "Adrian Chaparro Montoya", "Conchita Chaparro Campuzano",
    "Diana Chaparro Villalobos", "Dora Elda Chaparro campuzano", "Dolores Chaparro Campuzano",
    "Enedina Campuzano Gutierrez", "Eloy Eduardo Chaparro Villalobos", "Ivonne Chaparro Campuzano",
    "Jorge Moreno", "Liliana Melendez", "Lucio Perez", "Manuel Balderrama",
    "Maria del Rayo Montoya", "Marisol Barrera Chaparro", "Noe Chaparro Campuzano",
    "Ricardo Valenzuela Chaparro", "Rogelio Barrera", "Rogelio Barrera Chaparro",
    "Rosendo Chaparro Campuzano", "Teresa Chaparro Campuzano", "Teresa Barrera Chaparro"
]
def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # height in twentieths of a point (twips)
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def generar_docx():
    doc = Document()

    def add_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Página 1
    add_title("MINUTA FAF")
    doc.add_paragraph("\nSESIONAMOS BAJO EL SIGUIENTE ORDEN DEL DÍA:").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n1. Lectura del acta anterior")
    doc.add_paragraph("2. Lectura y revisión del estado financiero")
    doc.add_paragraph("3. Asuntos Generales")
    doc.add_paragraph("\nIniciamos nuestra reunión, dando lectura al acta de nuestra sesión anterior, habiendo sido aceptada y firmada de conformidad por los asistentes.")
    doc.add_paragraph("Continuando con el orden del día, se procede a dar lectura y revisión del estado financiero por nuestra tesorera Teresa Chaparro, destacando la siguiente información:")

    # Tabla de Ingresos
    tabla_ing = doc.add_table(rows=1, cols=2)
    tabla_ing.style = 'Table Grid'
    tabla_ing.rows[0].cells[0].text = "Concepto"
    tabla_ing.rows[0].cells[1].text = "Monto"

    for k, v in ingresos.items():
        row = tabla_ing.add_row().cells
        row[0].text = k
        row[1].text = f"${v:,.2f}"
        for paragraph in row[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_vertical_alignment(row[1], "center")

    doc.add_paragraph()  # Esto añade un salto de línea (br visual)

    # Tabla de Egresos
    tabla_egr = doc.add_table(rows=1, cols=2)
    tabla_egr.style = 'Table Grid'
    tabla_egr.rows[0].cells[0].text = "Concepto"
    tabla_egr.rows[0].cells[1].text = "Monto"

    for k, v in egresos.items():
        row = tabla_egr.add_row().cells
        row[0].text = k
        row[1].text = f"${v:,.2f}"
        for paragraph in row[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_vertical_alignment(row[1], "center")
    doc.add_paragraph(
        f"\nComo siguiente punto de la orden del día se llevó a cabo la cooperación correspondiente a {fecha}, "
        f"reuniéndose $2,500.00 que se depositarán en el fondo del FAF."
    )

    doc.add_paragraph(
        f"\nNo habiendo más asuntos que tratar se dio por terminada la sesión, siendo las 20:00 horas, "
        f"no sin antes agradecer el delicioso ambigú a cargo de {ambigu}."
    )

    doc.add_paragraph("\n\nROSENDO CHAPARRO CAMPUZANO\nSECRETARIO DEL FAF.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Página 2
    add_title("MINUTA FAF")
    doc.add_paragraph(
        f"\nSiendo las 18:00 horas del día {fecha} se llevó a cabo la sesión ordinaria del Fondo de Ahorro Familiar, "
        f"correspondiente a los meses de {meses}, contando con la presencia de los siguientes miembros:"
    )

    tabla_firmas = doc.add_table(rows=1, cols=2)
    tabla_firmas.style = 'Table Grid'
    tabla_firmas.autofit = False

    tabla_firmas.columns[0].width = Inches(3)
    tabla_firmas.columns[1].width = Inches(3.5)

    tabla_firmas.rows[0].cells[0].text = "Nombre"
    tabla_firmas.rows[0].cells[1].text = "Firma"

    for nombre in miembros:
        row = tabla_firmas.add_row()
        row.cells[0].text = nombre
        row.cells[1].text = ""  # espacio en blanco para firmar
        
        set_row_height(row, 0.7)  # Pasa el objeto fila correcto



    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

if st.button("📄 Generar Minuta"):
    archivo = generar_docx()
    st.success("✅ Documento generado con éxito.")
    st.download_button(
        label="⬇️ Descargar minuta en Word",
        data=archivo,
        file_name="minuta_faf.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
